import docx
import word

docName = 'sample-one-line.docx'
print('Document in full :\n',word.getTextWord(docName))

doc = docx.Document(docName)
print('단락 수 :',len(doc.paragraphs))
print('2번 단락 :',doc.paragraphs[1].text)
print('2번 단락 스타일 :',doc.paragraphs[1].style)
print('Paragraph 1:',doc.paragraphs[0].text)
print('Number of runs in paragraph 1:',len(doc.paragraphs[0].runs))
for idx, run in enumerate(doc.paragraphs[0].runs):
    print('Run %s : %s' %(idx,run.text))

print('is Run 0 underlined:',doc.paragraphs[0].runs[5].underline)
print('is Run 2 bold:',doc.paragraphs[0].runs[1].bold)
print('is Run 7 italic:',doc.paragraphs[0].runs[3].italic)